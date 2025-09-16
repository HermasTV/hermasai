import os
from docx import Document
from openai import OpenAI

# Initialize OpenAI client (make sure you set OPENAI_API_KEY in env)
client = OpenAI()

def extract_text_for_ai(docx_path):
    """
    Extracts plain text from the resume for AI processing.
    """
    doc = Document(docx_path)
    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append(p.text)
    return "\n".join(paragraphs)

def update_docx_with_ai(docx_in, docx_out, ai_text):
    """
    Replaces text inside the existing runs while preserving
    formatting, hyperlinks, and layout.
    """
    doc = Document(docx_in)
    ai_lines = ai_text.split("\n")
    para_idx = 0

    for p in doc.paragraphs:
        if para_idx >= len(ai_lines):
            break

        new_text = ai_lines[para_idx].strip()
        old_text = p.text.strip()

        if not new_text or old_text == "":
            para_idx += 1
            continue

        # Replace run text while preserving style
        remaining = new_text
        for r in p.runs:
            if remaining:
                r_len = len(r.text)
                r.text = remaining[:r_len] or ""
                remaining = remaining[r_len:]
            else:
                r.text = ""
        if remaining:  # add leftover text if AI made it longer
            p.add_run(remaining)

        para_idx += 1

    doc.save(docx_out)

def get_ai_corrected_text(text):
    """
    Sends resume text to OpenAI for grammar/style correction.
    """
    prompt = f"""
    You are a professional resume editor.
    Please correct grammar, spelling, and clarity in the following resume text
    while keeping the structure (headings, bullet points, job descriptions) intact.

    Resume:
    {text}
    """

    response = client.responses.create(
        model="gpt-4.1-mini",  # you can change to gpt-4.1 or gpt-4o
        input=prompt
    )

    return response.output_text.strip()

if __name__ == "__main__":
    resume_path = "/mnt/d/Projects/hermas.ai/hermasai/resources/resume-org.docx"           # Input resume file
    output_path = "resume_edited.docx"    # Output file

    # 1. Extract plain text for AI
    plain_text = extract_text_for_ai(resume_path)

    # 2. Get AI corrected version
    ai_fixed_text = get_ai_corrected_text(plain_text)

    # 3. Update docx while preserving formatting
    update_docx_with_ai(resume_path, output_path, ai_fixed_text)

    print(f"✅ Resume updated and saved at {output_path}")
